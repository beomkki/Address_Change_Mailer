#!/usr/bin/env python3
"""Generate grouped Outlook draft messages for address change notifications.

Groups trademarks by country and creates one email per country with a dynamic table
containing all trademarks for that country.
"""

from __future__ import annotations

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Dict, List

import openpyxl
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

try:
    from docx2msg import Docx2Msg
except ImportError as exc:
    raise SystemExit(
        "docx2msg 패키지를 찾을 수 없습니다. 먼저 `pip install docx2msg`를 실행해 주세요."
    ) from exc


def normalize_value(value) -> str:
    """Normalize Excel cell value to string."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value)
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    return str(value)


def load_marks_by_country(excel_path: Path) -> Dict[str, List[Dict[str, str]]]:
    """Load trademarks from Excel and group them by country code.

    Returns:
        Dictionary mapping country code to list of trademark records.
    """
    workbook = openpyxl.load_workbook(excel_path, data_only=True)
    worksheet = workbook.active

    rows_iter = worksheet.iter_rows(values_only=True)
    headers = [normalize_value(cell) or "" for cell in next(rows_iter)]

    grouped: Dict[str, List[Dict[str, str]]] = {}

    for values in rows_iter:
        row_dict: Dict[str, str] = {}
        empty = True
        for header, value in zip(headers, values):
            if not header:
                continue
            text_value = normalize_value(value)
            if text_value:
                empty = False
            row_dict[header] = text_value

        if not empty and row_dict.get("Country Code"):
            country_code = row_dict["Country Code"]
            if country_code not in grouped:
                grouped[country_code] = []
            grouped[country_code].append(row_dict)

    return grouped


def load_recipient_mapping(excel_path: Path) -> Dict[str, Dict[str, str]]:
    """Load recipient mapping from mailing list Excel file.

    Returns:
        Dictionary mapping country code to recipient info (To, CC, etc).
    """
    workbook = openpyxl.load_workbook(excel_path, data_only=True)

    mapping: Dict[str, Dict[str, str]] = {}

    # Try Sheet1 first (has headers)
    if "Sheet1" in workbook.sheetnames:
        ws = workbook["Sheet1"]
        rows_iter = ws.iter_rows(values_only=True)
        headers = next(rows_iter)  # Skip header row

        for row in rows_iter:
            if row[0]:  # Country code
                country_code = normalize_value(row[0])
                mapping[country_code] = {
                    "country_name": normalize_value(row[2]) if len(row) > 2 else "",
                    "from": normalize_value(row[3]) if len(row) > 3 else "",
                    "to": normalize_value(row[4]) if len(row) > 4 else "",
                    "cc": normalize_value(row[5]) if len(row) > 5 else "",
                }

    # Try Sheet2 (no headers)
    if "Sheet2" in workbook.sheetnames:
        ws = workbook["Sheet2"]
        rows_iter = ws.iter_rows(values_only=True)
        next(rows_iter, None)  # Skip first row if exists

        for row in rows_iter:
            if row[0]:  # Country code
                country_code = normalize_value(row[0])
                # Don't overwrite if already exists from Sheet1
                if country_code not in mapping:
                    mapping[country_code] = {
                        "country_name": normalize_value(row[1]) if len(row) > 1 else "",
                        "from": normalize_value(row[2]) if len(row) > 2 else "",
                        "to": normalize_value(row[3]) if len(row) > 3 else "",
                        "cc": normalize_value(row[4]) if len(row) > 4 else "",
                    }

    return mapping


def add_table_row(table, row_data: List[str]):
    """Add a new row to a Word table with the given data.

    Args:
        table: Word table object
        row_data: List of cell values for the new row
    """
    new_row = table.add_row()
    for idx, value in enumerate(row_data):
        if idx < len(new_row.cells):
            cell = new_row.cells[idx]
            cell.text = value
            # Set font size to match existing cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def prepare_template_with_marks(
    template_path: Path,
    marks: List[Dict[str, str]],
    tmp_dir: Path,
    country_name: str = "",
    subject: str = "",
) -> Path:
    """Prepare Word template by populating the first table with trademark data.

    Args:
        template_path: Path to the original Word template
        marks: List of trademark records to insert
        tmp_dir: Temporary directory for the modified template
        country_name: Name of the country (for context)
        subject: Email subject line

    Returns:
        Path to the modified template
    """
    doc = Document(template_path)

    if not doc.tables:
        raise ValueError("템플릿에 테이블이 없습니다.")

    # Add YAML header for docx2msg (required for subject line)
    header = doc.sections[0].header
    # Clear existing header paragraphs
    for para in list(header.paragraphs):
        header._element.remove(para._element)
    # Add YAML front matter (no closing --- to avoid multiple document error)
    # Escape double quotes in subject for YAML
    safe_subject = subject.replace('"', '\\"')
    header.add_paragraph("---")
    header.add_paragraph(f'Subject: "{safe_subject}"')

    # Get the first table (trademark table)
    table = doc.tables[0]

    # Clear existing data rows (keep header row 0)
    # Remove rows in reverse order to avoid index issues
    for _ in range(len(table.rows) - 1):
        table._element.remove(table.rows[-1]._element)

    # Add trademark data rows
    for mark in marks:
        row_data = [
            mark.get("Mark", ""),
            mark.get("Class", ""),
            mark.get("Appl. Date", ""),
            mark.get("Appl. No.", ""),
            mark.get("Reg. Date", ""),
            mark.get("Reg. No.", ""),
        ]
        add_table_row(table, row_data)

    # Save modified template
    temp_path = tmp_dir / f"template_{country_name}.docx"
    doc.save(temp_path)
    return temp_path


def sanitize_filename(preferred: str, alternate: str, index: int) -> str:
    """Create a safe filename from the given string."""
    candidate = preferred.strip() if preferred else ""
    candidate = re.sub(r"[^A-Za-z0-9._-]+", "_", candidate)
    candidate = candidate.strip("._")
    if not candidate:
        candidate = re.sub(r"[^A-Za-z0-9._-]+", "_", alternate).strip("._")
    if not candidate:
        candidate = f"message_{index:02d}"
    return candidate


def run_address_change_mail_merge(
    marks_excel: str | Path,
    mailing_list_excel: str | Path,
    template_path: str | Path,
    output_dir: str | Path,
) -> int:
    """Generate address change emails grouped by country.

    Args:
        marks_excel: Path to Excel file with trademark list
        mailing_list_excel: Path to Excel file with recipient mapping
        template_path: Path to Word template
        output_dir: Directory to save generated MSG files

    Returns:
        Number of MSG files generated
    """
    base_dir = Path(__file__).resolve().parent

    marks_excel = Path(marks_excel)
    if not marks_excel.is_absolute():
        marks_excel = base_dir / marks_excel

    mailing_list_excel = Path(mailing_list_excel)
    if not mailing_list_excel.is_absolute():
        mailing_list_excel = base_dir / mailing_list_excel

    template_path = Path(template_path)
    if not template_path.is_absolute():
        template_path = base_dir / template_path

    output_dir = Path(output_dir)
    if not output_dir.is_absolute():
        output_dir = base_dir / output_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    # Validate files exist
    if not marks_excel.exists():
        raise SystemExit(f"상표 리스트 파일을 찾을 수 없습니다: {marks_excel}")
    if not template_path.exists():
        raise SystemExit(f"워드 템플릿을 찾을 수 없습니다: {template_path}")

    # Load data
    print("Loading trademark data...")
    marks_by_country = load_marks_by_country(marks_excel)
    print(f"Found {len(marks_by_country)} countries with {sum(len(m) for m in marks_by_country.values())} total marks")

    # Load recipient mapping (optional)
    recipient_mapping = {}
    if mailing_list_excel.exists():
        print("Loading recipient mapping from mailing list...")
        recipient_mapping = load_recipient_mapping(mailing_list_excel)
        print(f"Found {len(recipient_mapping)} recipient mappings")
    else:
        print("Mailing list file not found - will use recipient info from marks file only")

    generated = 0

    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)

        for index, (country_code, marks) in enumerate(sorted(marks_by_country.items()), start=1):
            print(f"\nProcessing {country_code}: {len(marks)} marks")

            # Priority 1: Check if marks data has recipient info (수신, 참조, 국가명칭)
            # Use the first mark's data as representative for the country
            first_mark = marks[0] if marks else {}
            to_value = first_mark.get("수신", "") or first_mark.get("To", "")
            cc_value = first_mark.get("참조", "") or first_mark.get("CC", "")
            country_name = first_mark.get("국가명칭", "") or first_mark.get("Country Name", "") or first_mark.get("Country", "")
            our_ref = first_mark.get("Our. Ref.", "") or first_mark.get("참조번호", "") or "Our. Ref."

            # Priority 2: If not in marks data, lookup from recipient mapping
            if not to_value:
                recipient_info = recipient_mapping.get(country_code)
                if recipient_info:
                    to_value = recipient_info.get("to", "")
                    cc_value = cc_value or recipient_info.get("cc", "")
                    country_name = country_name or recipient_info.get("country_name", country_code)
                else:
                    print(f"  Warning: No recipient info found (neither in marks nor mailing list), skipping...")
                    continue

            # Fallback for country name
            if not country_name:
                country_name = country_code

            if not to_value:
                print(f"  Warning: No 'To' address for {country_code}, skipping...")
                continue

            print(f"  To: {to_value[:50]}{'...' if len(to_value) > 50 else ''}")
            if cc_value:
                print(f"  CC: {cc_value[:50]}{'...' if len(cc_value) > 50 else ''}")

            # Generate MSG file metadata
            subject = f"({our_ref})  Inquiry regarding Recordal of Change of Address ({country_name})"
            base_name = sanitize_filename(f"{country_code}_{country_name}_AddressChange", country_code, index)

            # Prepare template with marks
            prepared_template = prepare_template_with_marks(
                template_path, marks, tmp_path, country_code, subject
            )

            try:
                with Docx2Msg(prepared_template) as converter:
                    mail = converter.convert()
                    mail.Subject = subject
                    if to_value:
                        mail.To = to_value
                    if cc_value:
                        mail.CC = cc_value

                    msg_path = output_dir / f"{base_name}.msg"
                    mail.SaveAs(str(msg_path.resolve()), 3)
                    generated += 1
                    print(f"  ✓ Saved: {msg_path.name}")

                    mail.Close(False)
            except Exception as e:
                print(f"  ✗ Error generating MSG for {country_code}: {e}")
                continue

    print(f"\n완료: {generated}개의 MSG 파일을 {output_dir}에 생성했습니다.")
    return generated


def build_argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="주소변경 안내 메일을 국가별로 그룹핑하여 생성합니다.",
    )
    parser.add_argument(
        "--marks-excel",
        default="List of Marks.xlsx",
        help="상표 리스트 엑셀 파일 경로",
    )
    parser.add_argument(
        "--mailing-list",
        default="메일링 리스트.xlsx",
        help="수신인 매핑 엑셀 파일 경로",
    )
    parser.add_argument(
        "--template",
        default="Address_Change_Mail_Sample.docx",
        help="메일 본문 워드 템플릿 경로",
    )
    parser.add_argument(
        "--output-dir",
        default="output-address-change",
        help="생성된 MSG 파일을 저장할 디렉터리",
    )
    return parser


def main() -> None:
    parser = build_argument_parser()
    args = parser.parse_args()

    run_address_change_mail_merge(
        marks_excel=args.marks_excel,
        mailing_list_excel=args.mailing_list,
        template_path=args.template,
        output_dir=args.output_dir,
    )


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        sys.exit(1)
