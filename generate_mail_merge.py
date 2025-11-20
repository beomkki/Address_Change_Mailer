#!/usr/bin/env python3
"""Generate Outlook draft messages from Excel/Word mail-merge data using docx2msg."""

from __future__ import annotations

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Dict, Iterable, List

import openpyxl
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    from docx2msg import Docx2Msg
except ImportError as exc:  # pragma: no cover - dependency missing
    raise SystemExit(
        "docx2msg 패키지를 찾을 수 없습니다. 먼저 `pip install docx2msg`를 실행해 주세요."
    ) from exc

FIELD_PATTERN = re.compile(r"«([^»]+)»")
DEFAULT_SUBJECT_TEMPLATE = "(«관리번호»«국가코드») New trademark application(s) in «국가명칭»"
HEADER_YAML_LINES = (
    "---",
    "Subject: {{ subject }}",
)


def iter_block_items(doc: _Document) -> Iterable[Paragraph | Table]:
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def collect_fields(template_path: Path) -> List[str]:
    document = Document(template_path)
    fields: set[str] = set()
    for item in iter_block_items(document):
        if isinstance(item, Paragraph):
            fields.update(FIELD_PATTERN.findall(item.text))
        else:
            for row in item.rows:
                for cell in row.cells:
                    fields.update(FIELD_PATTERN.findall(cell.text))
    return sorted(fields)


def normalize_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int,)):
        return str(value)
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value)
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    return str(value)


def load_rows(excel_path: Path) -> List[Dict[str, str]]:
    workbook = openpyxl.load_workbook(excel_path, data_only=True)
    worksheet = workbook.active

    rows_iter = worksheet.iter_rows(values_only=True)
    headers = [normalize_value(cell) or "" for cell in next(rows_iter)]

    data_rows: List[Dict[str, str]] = []
    for values in rows_iter:
        row_dict: Dict[str, str] = {}
        empty = True
        for header, value in zip(headers, values):
            if not header:
                continue
            text_value = normalize_value(value)
            if text_value:
                empty = False
            row_dict[header] = text_value
        if not empty:
            data_rows.append(row_dict)

    return data_rows


def convert_paragraph_placeholders(paragraph: Paragraph) -> None:
    runs = list(paragraph.runs)
    i = 0
    while i < len(runs):
        run = runs[i]
        text = run.text
        if not text:
            i += 1
            continue
        if "«" in text and "»" in text:
            run.text = re.sub(r"«([^»]+)»", r"{{ \1 }}", text)
            i += 1
            continue
        if text == "«":
            start_run = run
            i += 1
            name_parts: List[str] = []
            while i < len(runs) and runs[i].text != "»":
                name_parts.append(runs[i].text)
                runs[i].text = ""
                i += 1
            field_name = "".join(name_parts)
            start_run.text = f"{{{{ {field_name} }}}}"
            if i < len(runs) and runs[i].text == "»":
                runs[i].text = ""
                i += 1
            continue
        if "«" in text:
            run.text = text.replace("«", "{{ ")
        if "»" in text:
            run.text = run.text.replace("»", " }}")
        i += 1


def prepare_template(original_template: Path, tmp_dir: Path) -> Path:
    document = Document(original_template)

    for item in iter_block_items(document):
        if isinstance(item, Paragraph):
            convert_paragraph_placeholders(item)
        else:
            for row in item.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        convert_paragraph_placeholders(para)

    header = document.sections[0].header
    # remove existing header paragraphs
    for para in list(header.paragraphs):
        header._element.remove(para._element)
    for line in HEADER_YAML_LINES:
        header.add_paragraph(line)

    temp_path = tmp_dir / "mail_template.docx"
    document.save(temp_path)
    return temp_path


def escape_docx_text(value: str) -> str:
    value = value.replace('&', '&amp;')
    value = value.replace('<', '&lt;')
    value = value.replace('>', '&gt;')
    return value


def sanitize_filename(preferred: str, alternate: str, index: int) -> str:
    candidate = preferred.strip() if preferred else ""
    candidate = re.sub(r"[^A-Za-z0-9._-]+", "_", candidate)
    candidate = candidate.strip("._")
    if not candidate:
        candidate = re.sub(r"[^A-Za-z0-9._-]+", "_", alternate).strip("._")
    if not candidate:
        candidate = f"message_{index:02d}"
    return candidate


def normalize_field_markers(text: str) -> str:
    """Ensure guillemet-style placeholders even when << >> is used."""
    return text.replace("<<", "«").replace(">>", "»")


def extract_template_fields(text: str) -> List[str]:
    """Extract field names from a template string with « » or << >> markers."""
    normalized = normalize_field_markers(text)
    return FIELD_PATTERN.findall(normalized)


def build_argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="엑셀/워드 병합 데이터를 활용해 docx2msg로 Outlook 초안을 생성합니다.",
    )
    parser.add_argument("--excel", default="Filing_Merge.xlsx", help="병합 데이터가 담긴 엑셀 파일 경로")
    parser.add_argument("--template", default="Filing.docx", help="메일 본문 워드 템플릿 경로")
    parser.add_argument("--output-dir", default="output-msg", help="생성된 MSG 파일을 저장할 디렉터리")
    parser.add_argument("--to-field", default="수신", help="엑셀에서 수신인(To)에 해당하는 컬럼명")
    parser.add_argument("--cc-field", default="참조", help="엑셀에서 참조(CC)에 해당하는 컬럼명")
    parser.add_argument(
        "--subject-template",
        default="",
        help="기본 제목 템플릿을 덮어쓸 문자열 (docx 본문에도 적용)",
    )
    parser.add_argument("--attachment-field", default="첨부파일", help="엑셀에서 첨부파일 경로에 해당하는 컬럼명 (세미콜론으로 구분)")
    return parser


def run_mail_merge(
    excel_path: str | Path,
    template_path: str | Path,
    output_dir: str | Path,
    *,
    to_field: str = "수신",
    cc_field: str = "참조",
    subject_template: str | None = "",
    attachment_field: str = "첨부파일",
) -> int:
    base_dir = Path(__file__).resolve().parent

    excel_path = Path(excel_path)
    if not excel_path.is_absolute():
        excel_path = base_dir / excel_path

    template_path = Path(template_path)
    if not template_path.is_absolute():
        template_path = base_dir / template_path

    output_dir = Path(output_dir)
    if not output_dir.is_absolute():
        output_dir = base_dir / output_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    if not excel_path.exists():
        raise SystemExit(f"엑셀 파일을 찾을 수 없습니다: {excel_path}")
    if not template_path.exists():
        raise SystemExit(f"워드 템플릿을 찾을 수 없습니다: {template_path}")

    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        prepared_template = prepare_template(template_path, tmp_path)
        available_fields = collect_fields(template_path)

        rows = load_rows(excel_path)
        if not rows:
            raise SystemExit("엑셀에서 데이터를 찾지 못했습니다.")

        subject_template = normalize_field_markers(subject_template or DEFAULT_SUBJECT_TEMPLATE)

        # subject_template에서 필드 추출하여 available_fields에 추가
        subject_fields = extract_template_fields(subject_template)
        all_fields = sorted(set(available_fields) | set(subject_fields))

        missing_fields: List[str] = []
        for field in all_fields:
            if all(field not in row or not row[field] for row in rows):
                missing_fields.append(field)
        if missing_fields:
            print(
                "경고: 아래 필드는 엑셀 데이터에서 빈 값입니다 → "
                + ", ".join(missing_fields)
            )

        generated = 0
        for index, row in enumerate(rows, start=1):
            raw_mapping = {key: row.get(key, "") for key in all_fields}
            escaped_mapping = {key: escape_docx_text(raw_mapping[key]) for key in available_fields}

            subject = subject_template
            for key, value in raw_mapping.items():
                subject = subject.replace(f"«{key}»", value)
                subject = subject.replace(f"<<{key}>>", value)

            to_value = row.get(to_field, "")
            cc_value = row.get(cc_field, "")

            base_name = sanitize_filename(subject, subject, index)

            context = {**escaped_mapping, "subject": escape_docx_text(subject)}

            with Docx2Msg(prepared_template) as converter:
                converter.template.render(context)
                mail = converter.convert()
                if to_value:
                    mail.To = to_value
                if cc_value:
                    mail.CC = cc_value
                html_body = mail.HTMLBody

                # 첨부파일 처리
                attachment_value = row.get(attachment_field, "")
                if attachment_value:
                    attachment_paths = [p.strip() for p in attachment_value.split(";") if p.strip()]
                    for attachment_path_str in attachment_paths:
                        attachment_path = Path(attachment_path_str)
                        if not attachment_path.is_absolute():
                            attachment_path = base_dir / attachment_path

                        if attachment_path.exists():
                            mail.Attachments.Add(str(attachment_path.resolve()))
                            print(f"  Added attachment: {attachment_path.name}")
                        else:
                            print(f"  Warning: Attachment not found: {attachment_path}")

                msg_path = output_dir / f"{base_name}.msg"
                mail.SaveAs(str(msg_path.resolve()), 3)
                generated += 1
                print(f"Saved MSG: {msg_path}")

                mail.Close(False)

        print(f"완료: {generated}개의 MSG 파일을 {output_dir}에 생성했습니다.")

    return generated

def main() -> None:
    parser = build_argument_parser()
    args = parser.parse_args()

    run_mail_merge(
        excel_path=args.excel,
        template_path=args.template,
        output_dir=args.output_dir,
        to_field=args.to_field,
        cc_field=args.cc_field,
        subject_template=args.subject_template,
        attachment_field=args.attachment_field,
    )


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        sys.exit(1)
